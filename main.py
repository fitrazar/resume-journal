import fitz
import google.generativeai as genai
from docx import Document

MODEL_NAME = "gemini-1.0-pro";
API_KEY = "GEMINI_API_KEY";


genai.configure(api_key=API_KEY)

# Set up the model
generation_config = {
  "temperature": 0.9,
  "top_p": 1,
  "top_k": 1,
  "max_output_tokens": 2048,
}

safety_settings = [
  {
    "category": "HARM_CATEGORY_HARASSMENT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
  },
  {
    "category": "HARM_CATEGORY_HATE_SPEECH",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
  },
  {
    "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
  },
  {
    "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
  },
]

model = genai.GenerativeModel(model_name=MODEL_NAME,
                              generation_config=generation_config,
                              safety_settings=safety_settings)

convo = model.start_chat(history=[
])

doc = fitz.open("file.pdf")

text = ""
for page in doc: 
    text += page.get_text()


def format_text_to_docx(generated):
    document = Document()

    lines = generated.split('\n')

    for line in lines:
        formatted_text, is_bold, is_italic = format_text(line)
        paragraph = document.add_paragraph()
        run = paragraph.add_run(formatted_text)
        if is_bold:
            run.bold = True
        elif is_italic:
            run.italic = True
        

    return document

def format_text(line):
    text_line = line.strip()
    
    if text_line.startswith("**") and text_line.endswith("**"):
        formatted_text = text_line[2:-2]
        is_bold = True
        is_italic = False
    elif text_line.startswith("*"):
        formatted_text = text_line[1:-1]
        is_bold = False
        is_italic = True
    else:
        formatted_text = text_line
        is_bold = False
        is_italic = False
    
    return formatted_text, is_bold, is_italic

  
def main():
    print("Pilih opsi:")
    print("1. AI")
    print("2. Resume")

    pilihan = input("Masukkan pilihan Anda (1 atau 2): ")

    if pilihan == "1":

        pertanyaan = input("Masukkan pertanyaan Anda: ")

        convo.send_message(text + '\n\n' +pertanyaan)

        print(convo.last.text)

    elif pilihan == "2":
        convo.send_message("Tolong buatkan saya resume tentang kalimat ini dengan lengkap dan jelas, serta berikan pendapat dan kesimpulan anda tentang apa yang dibahas pada kalimat tersebut \n" + text)

        generated = convo.last.text
        document = format_text_to_docx(generated)
        document.save('resume.docx')

        print("Resume telah disimpan ke file resume.docx")

    else:
        print("Pilihan tidak valid.")


if __name__ == "__main__":
    main()
