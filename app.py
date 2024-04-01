import fitz
import google.generativeai as genai
from docx import Document

MODEL_NAME = "gemini-1.0-pro";
API_KEY = "AIzaSyAYHYm8mF6kD0nMfb_GDVwRO6gZ-TR1Xv4";


doc = fitz.open("contoh2.pdf")
text = ""
for page in doc: 
  text += page.get_text()


genai.configure(api_key=API_KEY)

# Set up the model
generation_config = {
  "temperature": 0.9,
  "top_p": 1,
  "top_k": 1,
  "max_output_tokens": 2048,
}

safety_settings = [
  {
    "category": "HARM_CATEGORY_HARASSMENT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
  },
  {
    "category": "HARM_CATEGORY_HATE_SPEECH",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
  },
  {
    "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
  },
  {
    "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
  },
]

model = genai.GenerativeModel(model_name=MODEL_NAME,
                              generation_config=generation_config,
                              safety_settings=safety_settings)

convo = model.start_chat(history=[
])

convo.send_message("Tolong buatkan saya resume tentang kalimat ini dengan lengkap dan jelas, serta berikan pendapat dan kesimpulan anda tentang apa yang dibahas pada kalimat tersebut \n" + text)
print(convo.last.text)


def format_text(text):
    text = text.strip()
    
    if text.startswith("**") and text.endswith("**"):
        formatted_text = text[2:-2]
        is_bold = True
        is_italic = False
    elif text.startswith("*"):
        formatted_text = text[1:-1]
        is_bold = False
        is_italic = True
    else:
        formatted_text = text
        is_bold = False
        is_italic = False
    
    return formatted_text, is_bold, is_italic

document = Document()

generated = convo.last.text

lines = generated.split('\n')

for line in lines:
    formatted_text, is_bold, is_italic = format_text(line)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(formatted_text)
    if is_bold:
        run.bold = True
    elif is_italic:
        run.italic = True
        

document.save('resume.docx')